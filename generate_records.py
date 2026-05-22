import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

INPUT_FILE = "Records.md"
OUTPUT_FILE = "Records.docx"
TEMP_DIR = "temp_diagrams"

if not os.path.exists(TEMP_DIR):
    os.makedirs(TEMP_DIR)

def add_styled_paragraph(doc, text, bold=False, level=0, bullet=False):
    if level == 1:
        p = doc.add_heading(text, level=1)
    elif level == 2:
        p = doc.add_heading(text, level=2)
    elif level == 3:
        p = doc.add_heading(text, level=3)
    else:
        p = doc.add_paragraph(style='List Bullet' if bullet else 'Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
    return p

def generate():
    print("Reading Records.md...")
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # --- TITLE PAGE ---
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("LAB RECORD")
    run.bold = True
    run.font.size = Pt(24)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SYSTEM DESIGN & ANALYSIS")
    run.bold = True
    run.font.size = Pt(18)
    
    for _ in range(2): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Project: ReLaunchHer Platform").italic = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("May 2026")
    
    doc.add_page_break()

    # Split content into parts (text and mermaid)
    parts = re.split(r'(```mermaid[\s\S]*?```)', content)
    
    mermaid_count = 0
    for part in parts:
        if part.startswith("```mermaid"):
            mermaid_count += 1
            mermaid_code = part.replace("```mermaid", "").replace("```", "").strip()
            
            mm_file = os.path.join(TEMP_DIR, f"diagram_{mermaid_count}.mmd")
            img_file = os.path.join(TEMP_DIR, f"diagram_{mermaid_count}.png")
            
            with open(mm_file, "w", encoding="utf-8") as f:
                f.write(mermaid_code)
            
            print(f"Generating diagram {mermaid_count}...")
            try:
                # Use npx mmdc
                subprocess.run(["npx", "-y", "mmdc", "-i", mm_file, "-o", img_file, "-t", "forest", "-b", "transparent", "-w", "1200"], check=True, shell=True)
                doc.add_picture(img_file, width=Inches(5.5))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Error generating diagram {mermaid_count}: {e}")
                doc.add_paragraph(f"[Diagram Placeholder: {mermaid_count}]")
        else:
            lines = part.split("\n")
            for line in lines:
                line = line.strip()
                if not line: continue
                
                if line.startswith("# "):
                    add_styled_paragraph(doc, line[2:], level=1)
                elif line.startswith("## "):
                    add_styled_paragraph(doc, line[3:], level=2)
                elif line.startswith("### "):
                    add_styled_paragraph(doc, line[4:], level=3)
                elif line.startswith("- "):
                    add_styled_paragraph(doc, line[2:], bullet=True)
                elif line.startswith("**") and line.endswith("**"):
                    add_styled_paragraph(doc, line.replace("**", ""), bold=True)
                else:
                    add_styled_paragraph(doc, line)

    print(f"Saving {OUTPUT_FILE}...")
    doc.save(OUTPUT_FILE)
    print(f"✅ Success! {OUTPUT_FILE} generated.")

if __name__ == "__main__":
    generate()
