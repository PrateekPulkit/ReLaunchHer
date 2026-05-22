import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_styled_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def generate():
    doc = Document()
    
    # Global Style Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    team = [
        "Prateek Pulkit (AP23110011175)",
        "Srinadh Yakasiri (AP23110011171)",
        "Abhishek Das (AP23110011180)"
    ]

    # --- TITLE PAGE ---
    for _ in range(5): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("A REPORT ON").bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\nRELAUNCHHER: AN AI-DRIVEN CAREER RE-ENTRY PLATFORM FOR WOMEN WITH MULTI-FACTOR INFERENCE AND MULTILINGUAL ACCESSIBILITY")
    run.bold = True
    run.font.size = Pt(16)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\nBy\n")
    for member in team:
        doc.add_paragraph(member).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    for _ in range(3): doc.add_paragraph()
    doc.add_paragraph("Prepared in the partial fulfillment of the Project Based Learning of Course").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("CSE 306 – Software Engineering and Project Management").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(5): doc.add_paragraph()
    doc.add_paragraph("SRM UNIVERSITY, AP").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("April-May 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # --- ACKNOWLEDGEMENTS ---
    doc.add_heading("Acknowledgements", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "The successful completion of this project, ReLaunchHer, would not have been possible without the support and guidance of several individuals. I am deeply indebted to the following for their contribution to this work.")
    add_styled_paragraph(doc, "Firstly, I would like to express my sincere gratitude to the Vice Chancellor and Dean of SRM University, AP, for their vision in integrating Project-Based Learning into the curriculum and providing the state-of-the-art infrastructure required for this development. Their constant encouragement to solve real-world social problems through technology has been the primary inspiration for this project.")
    add_styled_paragraph(doc, "I am profoundly grateful to my Industry Mentor, whose professional insights into the modern hiring landscape and recruitment bottlenecks were instrumental in shaping the Readiness Index logic. Their real-world feedback helped bridge the gap between academic theory and industry application, ensuring that the platform's outputs are aligned with what recruiters actually value in returnees.")
    add_styled_paragraph(doc, "I would like to extend my deepest appreciation to DR. [Insert Faculty Name], Faculty Mentor in the Department of Computer Science and Engineering, for their continuous mentorship, technical guidance, and critical reviews. Their expertise in Software Engineering principles helped us refine our architectural decisions, UML modeling, and testing strategies. They pushed us to look beyond just the code and focus on the complete software lifecycle.")
    add_styled_paragraph(doc, "Finally, I want to thank my family and friends for their constant encouragement and patience during the long hours of coding and documentation. This project is a testament to the collaborative ecosystem fostered at SRM University, AP.")
    
    doc.add_page_break()

    # --- CERTIFICATE ---
    doc.add_heading("Certificate", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, f"This is to certify that the project entitled \"ReLaunchHer: An AI-Driven Career Re-Entry Platform for Women\" has been successfully completed by Prateek Pulkit (AP23110011175), Srinadh Yakasiri (AP23110011171), and Abhishek Das (AP23110011180) as a required academic component of the course CSE 306: Software Engineering and Project Management during Semester 6 of Academic Year 2025-26 in the Department of Computer Science and Engineering at SRM University, AP.")
    add_styled_paragraph(doc, "This work was executed under the guidance of the undersigned and is deemed to have successfully met all course requirements, including requirements specification, system modeling, implementation, and quality assurance testing as of May 2026.")
    for _ in range(5): doc.add_paragraph()
    doc.add_paragraph("(Signature of the Course Instructor)").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("DR. [Insert Name]").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("Department of CS and Engineering").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_page_break()

    # --- ABSTRACT ---
    doc.add_heading("Abstract", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "ReLaunchHer V1.0 is a comprehensive, multi-layered ecosystem designed to address the socio-economic challenge of career re-entry for women after professional breaks. In the current employment landscape, \"career gaps\" are often treated as liabilities, leading to an \"Invisibility Crisis\" where talented women are filtered out by rigid recruitment algorithms. This project proposes a Weighted Multi-Factor Inference Engine that quantifies professional potential by analyzing non-linear career trajectories, upskilling efforts, and life-skills developed during breaks.")
    add_styled_paragraph(doc, "Built with Node.js and SQLite, the platform follows an Iterative Enhancement Model of the SDLC. It features a high-fidelity, multilingual user interface supporting seven major Indian languages, ensuring inclusivity. The core innovation lies in the automated generation of 30/60/90-day personalized roadmaps that transform career anxiety into structured action. This report documents the complete engineering lifecycle—from SRS and feasibility analysis to UML modeling and algorithm complexity—demonstrating a robust solution for a high-impact social problem.")
    
    doc.add_page_break()

    # --- TABLE OF CONTENTS ---
    doc.add_heading("Table of Contents", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = [
        ("1. Introduction", "6"),
        ("2. Background Study & Literature Survey", "8"),
        ("3. Problem Statement & Feasibility Analysis", "10"),
        ("4. Proposed System Architecture", "12"),
        ("5. Software Requirements Specification (SRS)", "14"),
        ("6. Algorithms and Core Logic Modules", "16"),
        ("7. System Design and UML Modeling", "19"),
        ("8. Implementation & Project Management", "22"),
        ("9. Testing and Quality Assurance", "24"),
        ("10. Result / Output", "26"),
        ("11. Conclusions and Future Scope", "28"),
        ("12. References", "30")
    ]
    for text, page in toc:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text + " " + "." * (80 - len(text)) + " " + page)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    
    doc.add_page_break()

    # --- CHAPTER 1: INTRODUCTION ---
    doc.add_heading("1. Introduction", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "1.1 Project Overview", bold=True)
    add_styled_paragraph(doc, "ReLaunchHer is an AI-driven ecosystem designed to solve the systemic barriers women face when returning to the workforce. Career breaks, often necessitated by childcare, elderly care, or personal health, frequently lead to a \"career cliff\" rather than a temporary pause. ReLaunchHer acts as a digital career coach, utilizing data-driven inference to provide returnees with a clear, technical, and psychological roadmap back to professional success. The system is built with a focus on empathy, data accuracy, and scalability, ensuring that every user receives a personalized path tailored to their specific career history and target goals.")
    add_styled_paragraph(doc, "1.2 Motivation", bold=True)
    add_styled_paragraph(doc, "The motivation for this project stems from the documented \"Motherhood Penalty,\" where women's career trajectories are disproportionately affected by time away from work. In India, thousands of women with high-level technical expertise leave the workforce every year and find it nearly impossible to return due to algorithmic bias in standard job portals. ReLaunchHer aims to reclaim this lost economic potential by reframing the \"gap\" as a period of transition that can be managed through structured upskilling and confidence building. This project is a response to the need for a more empathetic and data-accurate recruitment pipeline that values resilience and life-skills as much as technical certifications.")
    add_styled_paragraph(doc, "1.3 Objectives of the Project", bold=True)
    add_styled_paragraph(doc, "Objective 1: To develop a multi-factor logic engine that quantifies professional readiness across technical and behavioral dimensions. This involves creating a scoring model that doesn't just look at years of experience, but also considers the depth of that experience and the currency of the skills involved.")
    add_styled_paragraph(doc, "Objective 2: To create an automated 30/60/90-day roadmap generator that provides actionable steps for skill gap closure. The roadmap must be dynamic, adjusting its milestones based on the user's available time and the urgency of their return.")
    add_styled_paragraph(doc, "Objective 3: To implement a multilingual interface that ensures accessibility for women from diverse linguistic backgrounds across India. By removing the language barrier, we aim to reach the 'last mile' of talent that is often excluded from the mainstream tech economy.")
    add_styled_paragraph(doc, "Objective 4: To build a role-based ecosystem connecting Returnees with Mentors and Recruiters in a unified pipeline, fostering a community of support.")
    
    doc.add_page_break()

    # --- CHAPTER 2: BACKGROUND STUDY ---
    doc.add_heading("2. Background Study & Literature Survey", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "2.1 Existing Systems Analysis", bold=True)
    add_styled_paragraph(doc, "Currently, the recruitment market is dominated by platforms like LinkedIn and Indeed. While these are excellent for continuous professional growth, they fail the \"non-linear\" user. Their search algorithms prioritize \"Recent Experience,\" which automatically pushes returnees to the bottom of the list. Furthermore, corporate returnship programs (like those at Amazon or Google) are highly competitive and only reach a tiny fraction of the returning population. There is a significant gap in the market for a platform that specifically addresses the needs of the mid-career returnee who requires more than just a job board—they require a transition framework. ReLaunchHer seeks to democratize this returnship logic by making high-quality career coaching logic accessible to everyone.")
    add_styled_paragraph(doc, "2.2 Literature Review", bold=True)
    add_styled_paragraph(doc, "Research by Correll et al. (2007) on the \"Motherhood Penalty\" provides the sociological foundation for this project, highlighting how the perception of \"competence\" drops when a gap appears on a resume. Technical research in Decision Support Systems (DSS) has shown that automated coaching tools can significantly reduce \"Decision Fatigue\" in career transitions. By automating the analysis of skill gaps, we can help returnees focus their limited time on the most impactful upskilling tasks. Furthermore, studies in Internationalization (i18n) emphasize that localized software significantly increases adoption rates in non-English speaking demographics, which is a core pillar of the ReLaunchHer mission.")
    add_styled_paragraph(doc, "2.3 Technological Context", bold=True)
    add_styled_paragraph(doc, "The project utilizes modern full-stack development patterns. Node.js provides a high-performance, non-blocking environment for the backend, while SQLite offers a robust, zero-configuration database solution for local persistence. The use of Vanilla JavaScript for the logic engine ensures that the system is fast, lightweight, and easy to maintain without the overhead of heavy frameworks.")
    
    doc.add_page_break()

    # --- CHAPTER 3: PROBLEM STATEMENT ---
    doc.add_heading("3. Problem Statement & Feasibility Analysis", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "3.1 Detailed Problem Statement", bold=True)
    add_styled_paragraph(doc, "The primary problem is the \"Algorithmic Invisibility\" of returnees in the standard hiring pipeline. This invisibility is caused by three core factors: (1) Absence of transferable skill mapping—traditional systems don't recognize the crisis management or multi-tasking developed during a break. (2) Lack of confidence-building mechanisms—the psychological toll of a career break is often ignored, leading to high dropout rates during the application process. (3) Language barriers—most career guidance is in English, excluding a large demographic of skilled women in India. Without a dedicated platform, the 'return to work' becomes a daunting, unmanaged project that leads to a significant loss of economic potential for society.")
    add_styled_paragraph(doc, "3.2 Feasibility Analysis", bold=True)
    add_styled_paragraph(doc, "3.2.1 Technical Feasibility", bold=True)
    add_styled_paragraph(doc, "The project utilizes Node.js and SQLite, which are mature, well-documented technologies. The logic engine is implemented via a modular JavaScript runtime, ensuring it can run on any modern web server without the need for expensive high-compute AI hardware. The use of CSS variables and modern layout techniques (Flexbox/Grid) ensures that the UI is responsive and accessible across devices, making the platform technically sound and deployable.")
    add_styled_paragraph(doc, "3.2.2 Economic Feasibility", bold=True)
    add_styled_paragraph(doc, "Since the platform uses a 'Privacy-First' local database approach (SQLite) and open-source technologies, the infrastructure costs are minimal. There are no expensive API dependencies or high-cost cloud computing requirements. This makes the platform extremely affordable to scale and maintain, ensuring high economic feasibility for widespread social deployment.")
    add_styled_paragraph(doc, "3.2.3 Operational Feasibility", bold=True)
    add_styled_paragraph(doc, "The role-based design (Returnee, Mentor, Recruiter) mirrors existing human workflows in the recruitment and coaching industry, making it easy for users to adopt. The intuitive glassmorphic UI reduces the learning curve, and the automated roadmapping reduces the operational load on human mentors, ensuring that the system can function effectively in real-world scenarios.")

    doc.add_page_break()

    # --- CHAPTER 4: PROPOSED SYSTEM ---
    doc.add_heading("4. Proposed System Architecture", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "4.1 Architectural Overview", bold=True)
    add_styled_paragraph(doc, "ReLaunchHer is architected as a Three-Tier Web Application to ensure scalability, security, and separation of concerns. The Presentation Tier is built using responsive HTML5 and CSS3, utilizing a modern design system with glassmorphic elements. The Logic Tier is managed by a Node.js server, which handles authentication, role-based routing, and the core inference logic. The Data Tier utilizes an SQLite database for persistent storage of user profiles, career plans, and job mappings. This architecture ensures that the system is modular, allowing for independent updates to each layer without affecting the others.")
    add_styled_paragraph(doc, "4.2 System Modules", bold=True)
    add_styled_paragraph(doc, "Module 1: Auth & Role Manager. This module handles secure user registration and login, utilizing hashed passwords and session-based authentication. It ensures that users are directed to the correct dashboard based on their role (Returnee, Mentor, or Recruiter).")
    add_styled_paragraph(doc, "Module 2: Inference Engine. The 'brain' of the platform, this module executes the weighted multi-factor scoring algorithm. It takes the user's profile data as input and produces a Readiness Index, which forms the basis for all further recommendations.")
    add_styled_paragraph(doc, "Module 3: Roadmap Synthesizer. This module converts the numerical scores and skill gaps into a structured, phase-based career plan. It utilizes a JSON-based template system to generate personalized milestones for 30, 60, and 90-day intervals.")
    add_styled_paragraph(doc, "Module 4: I18n Engine. A localized translation layer that dynamically swaps UI strings. It uses a key-value mapping system to provide real-time translation in 7+ languages, ensuring that the platform is accessible to a wider audience.")

    doc.add_page_break()

    # --- CHAPTER 5: SRS ---
    doc.add_heading("5. Software Requirements Specification (SRS)", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "5.1 Functional Requirements", bold=True)
    add_styled_paragraph(doc, "FR1: Role Registration. The system shall allow users to register as a Returnee, Mentor, or Recruiter with distinct profile fields for each role.")
    add_styled_paragraph(doc, "FR2: Career Assessment. The system shall provide a multi-step wizard for Returnees to input their career history, skills, and goals.")
    add_styled_paragraph(doc, "FR3: Readiness Scoring. The system shall calculate a Readiness Score (0-100) based on weighted factors including experience, skill match, and gap length.")
    add_styled_paragraph(doc, "FR4: Roadmap Generation. The system shall automatically generate a 30/60/90-day action plan for Returnees within 2 seconds of assessment completion.")
    add_styled_paragraph(doc, "FR5: Language Switching. The system shall allow users to switch the UI language between 7 Indian languages without requiring a page reload.")
    add_styled_paragraph(doc, "5.2 Non-Functional Requirements", bold=True)
    add_styled_paragraph(doc, "NFR1: Performance. The system shall have a page response time of less than 1.5 seconds on a standard 4G connection.")
    add_styled_paragraph(doc, "NFR2: Security. All user passwords must be hashed using a secure algorithm (e.g., bcrypt) before being stored in the database. All sessions must be managed via secure cookies.")
    add_styled_paragraph(doc, "NFR3: Scalability. The backend architecture shall support up to 100 concurrent database operations without performance degradation.")
    add_styled_paragraph(doc, "NFR4: Usability. The UI must follow a mobile-first responsive design, ensuring it is usable on devices with varying screen sizes.")

    doc.add_page_break()

    # --- CHAPTER 6: ALGORITHMS ---
    doc.add_heading("6. Algorithms and Core Logic Modules", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "6.1 Weighted Multi-Factor Scoring (Module 2)", bold=True)
    add_styled_paragraph(doc, "Theory: The Readiness Index is not a simple linear average. It utilizes a non-linear decay function for career gaps, recognizing that the first year of a gap has a different impact than the fifth. However, this is offset by an 'Active Upskilling' multiplier, where the penalty is reduced if the user demonstrates intentional learning during the break. The algorithm weights are carefully balanced: Experience (20%), Skill Match (22%), Confidence (12%), and Engagement (10%).")
    add_styled_paragraph(doc, "Algorithm Steps:", bold=True)
    add_styled_paragraph(doc, "1. Input Acquisition: Collect profile data (Exp, Gap, Skills, Role).")
    add_styled_paragraph(doc, "2. Base Score Calculation: Assign points based on years of previous experience (e.g., 2.5 pts per year).")
    add_styled_paragraph(doc, "3. Gap Penalty Application: Calculate penalty using a non-linear function: Penalty = (Months ^ 0.7) * 0.4.")
    add_styled_paragraph(doc, "4. Mitigation Analysis: Check reason for break. If reason is 'Education' or 'Upskilling', reduce penalty by 50%.")
    add_styled_paragraph(doc, "5. Skill Matching: Compare user skills against the target role's required skill catalog. Calculate match ratio.")
    add_styled_paragraph(doc, "6. Normalization and Clamping: Sum all scores and ensure the result is between 10 and 98.")
    add_styled_paragraph(doc, "6.2 Complexity Analysis", bold=True)
    add_styled_paragraph(doc, "Time Complexity: O(N * M) where N is the number of user skills and M is the market skill requirements for a specific role. This is highly efficient for real-time calculation.")
    add_styled_paragraph(doc, "Space Complexity: O(1) as the inference is performed in memory without requiring large temporary storage.")

    doc.add_page_break()

    # --- CHAPTER 7: SYSTEM DESIGN ---
    doc.add_heading("7. System Design and UML Modeling", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "The system design follows industry-standard modeling techniques to ensure structural integrity and clear data flow. This chapter outlines the logical and physical design of the ReLaunchHer platform.")
    add_styled_paragraph(doc, "7.1 Data Flow Diagram (DFD)", bold=True)
    add_styled_paragraph(doc, "Level 0 DFD: At the highest level, the system interacts with three primary entities: the Returnee, the Mentor, and the Recruiter. The Returnee provides profile data and receives a career plan. The Recruiter provides job postings and receives applicant lists.")
    add_styled_paragraph(doc, "Level 1 DFD: Within the system, data flows from the UI controller into the Auth module for validation, then into the Inference Engine for processing, and finally into the SQLite database for persistence. The results are then fetched by the Dashboard module to present to the user.")
    add_styled_paragraph(doc, "7.2 UML Class Diagram", bold=True)
    add_styled_paragraph(doc, "The class diagram shows the inheritance relationship where 'Returnee', 'Mentor', and 'Recruiter' all inherit from a base 'User' class. The 'Returnee' class has a one-to-one association with the 'CareerPlan' class, which is generated by the 'InferenceEngine'. This modular design allows for easy addition of new user types in the future.")
    add_styled_paragraph(doc, "7.3 Sequence Diagram", bold=True)
    add_styled_paragraph(doc, "The sequence diagram visualizes the asynchronous interaction during the career assessment. The UI controller sends a request to the server, which triggers the logic engine. The engine calculates the score and stores the result in the database before sending a confirmation back to the UI for redirection.")

    doc.add_page_break()

    # --- CHAPTER 8: IMPLEMENTATION ---
    doc.add_heading("8. Implementation & Project Management", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "8.1 Implementation Strategy", bold=True)
    add_styled_paragraph(doc, "The implementation followed a modular approach using a modern Node.js stack. To maximize performance and minimize bundle size, we utilized Vanilla JavaScript for all logic components, avoiding heavy client-side frameworks. For the UI, we implemented a custom CSS design system utilizing CSS Variables to manage our glassmorphic theme across the entire platform. The database was implemented using 'better-sqlite3', which provides a synchronous API for SQLite that is optimized for performance in Node.js environments.")
    add_styled_paragraph(doc, "8.2 SDLC Methodology: Iterative Enhancement", bold=True)
    add_styled_paragraph(doc, "The project followed an Iterative Enhancement Model. This allowed us to develop and test core functionalities (like the scoring logic) early in the process. Each iteration added a new layer of functionality: Iteration 1 focused on Auth and DB schema; Iteration 2 on the Inference Engine; and Iteration 3 on the Multilingual UI and Dashboard polish. This approach ensured that we always had a working prototype and could pivot quickly based on testing results.")
    add_styled_paragraph(doc, "8.3 Risk Management", bold=True)
    add_styled_paragraph(doc, "We identified and mitigated several risks during the development: (1) Data Security Risk—handled via server-side session management and role-based access control. (2) Algorithmic Bias Risk—mitigated by tuning the weighting factors to prioritize recent upskilling over historical gaps. (3) Performance Risk—handled by using a lightweight tech stack and optimizing database queries with proper indexing.")

    doc.add_page_break()

    # --- CHAPTER 9: TESTING ---
    doc.add_heading("9. Testing and Quality Assurance", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "Testing was integrated into every phase of the development lifecycle to ensure a high-quality, bug-free product. We utilized a combination of automated and manual testing strategies.")
    add_styled_paragraph(doc, "9.1 Unit Testing", bold=True)
    add_styled_paragraph(doc, "We developed an automated test suite using the Node.js native test runner to verify the core logic of the Inference Engine. We tested the 'scoreProfile' function against 20+ synthetic personas, including edge cases like 10-year gaps or zero previous experience, to ensure the output scores remained within the valid 10-98 range.")
    add_styled_paragraph(doc, "9.2 Integration Testing", bold=True)
    add_styled_paragraph(doc, "Integration tests were performed to ensure that the data flowed correctly between the UI, the Server, and the Database. We verified that a completed assessment correctly populated the 'plans' table in SQLite and that the Dashboard could successfully fetch and render this data.")
    add_styled_paragraph(doc, "9.3 User Acceptance Testing (UAT)", bold=True)
    add_styled_paragraph(doc, "We conducted a small-scale UAT with a cohort of women returning to the workforce. Their feedback was used to refine the language of the 90-day roadmaps, making them more encouraging and actionable. This ensured that the platform actually meets the psychological and technical needs of its primary users.")

    doc.add_page_break()

    # --- CHAPTER 10: RESULT ---
    doc.add_heading("10. Result / Output", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "The principal outcomes of the ReLaunchHer project demonstrate a successful application of software engineering principles to a complex social problem. The results show that automated career coaching is both feasible and highly effective.")
    add_styled_paragraph(doc, "10.1 Key Outcomes", bold=True)
    add_styled_paragraph(doc, "Outcome 1: Quantified Readiness. The platform successfully translates a complex, non-linear career history into a standardized metric (the Readiness Index), which helps both the candidate and the recruiter understand the candidate's potential.")
    add_styled_paragraph(doc, "Outcome 2: Dynamic Action Plans. Every user now receives a personalized 30/60/90-day roadmap that provides a clear path forward, significantly reducing the 'choice paralysis' often felt by returnees.")
    add_styled_paragraph(doc, "Outcome 3: Multilingual Inclusivity. The platform's ability to switch between 7 languages ensures that high-quality career guidance is accessible to women across India, regardless of their primary language.")
    
    doc.add_page_break()
    
    add_styled_paragraph(doc, "10.2 Visual Gallery of Outputs", bold=True)
    # Add screenshots here if they exist, else placeholders
    screenshot_files = [
        "screenshot_landing_1777485635274.png",
        "screenshot_hindi_1777485670623.png",
        "screenshot_planner_results_1777485835306.png",
        "screenshot_dashboard_1777485865834.png"
    ]
    for i, file in enumerate(screenshot_files):
        try:
            doc.add_picture(f"./screenshots/{file}", width=Inches(4.5))
            doc.add_paragraph(f"Figure {i+1}: ReLaunchHer platform output screenshot").alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            add_styled_paragraph(doc, f"[Image Placeholder: {file}]")
        doc.add_paragraph()

    doc.add_page_break()

    # --- CHAPTER 11: CONCLUSION ---
    doc.add_heading("11. Conclusions and Future Scope", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_styled_paragraph(doc, "11.1 Conclusion", bold=True)
    add_styled_paragraph(doc, "ReLaunchHer V1.0 successfully demonstrates that technology can be a powerful tool for social equity. By applying Software Engineering principles—specifically modular design, data-driven inference, and internationalization—the project transforms a daunting life transition into a manageable, structured project. The platform proves that the 'career gap' is not a liability, but a transition period that can be quantified and managed. The modularity of the system ensures that it can be easily updated as market trends and hiring standards evolve.")
    add_styled_paragraph(doc, "11.2 Future Scope", bold=True)
    add_styled_paragraph(doc, "AI Resume Parser: In future versions, we aim to implement a PDF analysis module that can automatically populate the career assessment by parsing uploaded resumes. This will significantly reduce the friction for new users.")
    add_styled_paragraph(doc, "Real-time Mentorship: We plan to integrate WebRTC-based video conferencing to allow Returnees and Mentors to connect for real-time interview practice and guidance directly within the platform.")
    add_styled_paragraph(doc, "Mobile Application: Developing a native mobile application using React Native will further increase accessibility for women in rural areas who primarily use mobile devices for internet access.")

    doc.add_page_break()

    # --- CHAPTER 12: REFERENCES ---
    doc.add_heading("12. References", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    refs = [
        "Correll, S. J., 'Getting a Job: Is There a Motherhood Penalty?' Amer. Jour. Soc., 112 (2007), pp. 1297-1338.",
        "Pressman, R. S., Software Engineering: A Practitioner's Approach 8th edition, McGraw-Hill (2014).",
        "iRelaunch Official Website, www.irelaunch.com (2026).",
        "Node.js Documentation, nodejs.org (2026).",
        "Better-SQLite3 Documentation, github.com/WiseLibs/better-sqlite3 (2026)."
    ]
    for i, ref in enumerate(refs):
        add_styled_paragraph(doc, f"{i+1}. {ref}")

    doc.save("report_v3_final.docx")
    print("report_v3_final.docx generated successfully with expanded content.")

if __name__ == "__main__":
    generate()
